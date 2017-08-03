from math import ceil
import re
import time 

import win32com.client
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

CH_2 = chr(0xb2)  #символ второй степени
CH_3 = chr(0xb3)  #символ третей степени
CH_GRAD = chr(0xb0)  #символ градуса
CH_D = chr(0xf8) #символ диаметра

CH_F = "$$$" #признак форматирования в строке: Ar Al Ac - выравнивание; B - жирный; U - подчеркнутый
#пример: "$$$Al;B;U$$$"

EXCEL_MAX_COL = 16384
EXCEL_MAX_ROW = 1048576

class Excel():
    """ Main Excel Automation object
    """
    def __init__(self, visible=False, fname=None, work_sheet=None):
        win32com.client.gencache.EnsureDispatch('Excel.Application')
        self.exapp = win32com.client.Dispatch("Excel.Application")  #переменная приложения Excel
        time.sleep(5)
        if visible:
            self.exapp.Visible = 1

        if fname is not None:
            time.sleep(5)
            self.wb = self.exapp.Workbooks.Open(fname) # переменная с текущим файлом эксель
        else:
            self.wb = None

        if work_sheet:
            time.sleep(5)
            self.sheet = self.exapp.Worksheets(work_sheet)
            self.sheet.Activate()
        else:
            self.sheet = self.exapp.ActiveSheet  #переменная текущего листа

    def close(self):
        # закрыть книгу, если открывали ее
        if self.wb:
            self.wb.Close()

    def quit(self):
        #закрыть эксель
        self.exapp.Quit()

    def select_sheet(self, sheet_name):
        #сменить текущий лист
        self.sheet = self.exapp.Worksheets(sheet_name)
        self.sheet.Activate()
        return self.sheet

    def get_row(self, row=1, first_col=1, empty_col=1):
        """ Возвращает в виде списка заданный ряд, начиная с first_col, и заканчивая после empty_col пустых ячеек
        """
        empty_cells = 0
        l=[]
        for i in range(first_col, EXCEL_MAX_COL+1):
            tag = self.sheet.Cells(row, i).Value
            #прекращаем скнировать после empty_col пустых ячеек подряд
            if tag:
                empty_cells = 0
                l.append(tag)
            else:
                empty_cells += 1
                if empty_cells >= empty_col:
                    break
        return l

    def get_range_data(self, r1, c1, r2, c2):
        #читает область ячеек в двумерный массив
        return self.sheet.Range(self.sheet.Cells(r1,c1), self.sheet.Cells(r2,c2)).Value

    def get_last_col(self, row):
        #возвращает последнюю непустую ячейку в строке 
        self.sheet.Cells(row, EXCEL_MAX_COL).Select()
        self.exapp.Selection.End(win32com.client.constants.xlToLeft).Select()
        return self.exapp.Selection.Column

    def get_last_row(self, col):
        #возвращает последнюю непустую ячейку в столбце 
        self.sheet.Cells(EXCEL_MAX_ROW, col).Select()
        self.exapp.Selection.End(win32com.client.constants.xlUp).Select()
        return self.exapp.Selection.Row


def word_new_row(wordapp,l):
    """ Add a row in Word table and fill it with data from l"""
    wordapp.Selection.InsertRowsBelow(1)
    for i in l:
        wordapp.Selection.TypeText(i)    
        wordapp.Selection.MoveRight(Unit=win32com.client.constants.wdCell)        


def GetCurrTableNum(wordapp):
#найти номер выбранной таблицы в документе или -1 если таблица не выбрана
    if not wordapp.Selection.Information(win32com.client.constants.wdWithInTable):
        return -1
    worddoc = wordapp.ActiveDocument
    for i in range(1, worddoc.Tables.Count+1):
        if (wordapp.Selection.Range.Start >= worddoc.Tables(i).Range.Start) and (wordapp.Selection.Range.End <= worddoc.Tables(i).Range.End):
            break
    return i            
        

def word_line_to_table_format_fast(table, row, l):
    """ Add a row in Word table and fill it with data from l, with formatting """
    table.add_row()


def word_line_to_table_format(table, row, l):
    """ Add a row in Word table and fill it with data from l, with formatting """
    table.add_row()
    for col, i in enumerate(l):
        rez=["",""]
        if CH_F in i:
            rez = i.split(CH_F)
            table.cell(row, col).text = rez[2] 
            for j in rez[1].split(";"):
                if j.upper()=="AL":
                    table.cell(row, col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif j.upper()=="AR":
                    table.cell(row, col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif j.upper()=="AC":
                    table.cell(row, col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif j.upper()=="B":
                    table.cell(row, col).paragraphs[0].runs[0].bold = True
                elif j.upper()=="U":
                    table.cell(row, col).paragraphs[0].runs[0].underline = True
        else:
            table.cell(row, col).text = i 
        #Вторую колонку по умолчанию сдвигаем влево
        if col==1 and not re.search("A[LRC]", rez[1], flags=re.IGNORECASE):
            table.cell(row, col).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT


def word_new_row_format(wordapp, l):
    """ Add a row in Word table and fill it with data from l, with formatting """
    wordapp.Selection.InsertRowsBelow(1)
    wordapp.Selection.Font.Bold = False
    wordapp.Selection.Font.Underline = win32com.client.constants.wdUnderlineNone
    wordapp.Selection.ParagraphFormat.Alignment = win32com.client.constants.wdAlignParagraphLeft
    for i in l:
        if CH_F in i:
            rez = i.split(CH_F)
            wordapp.Selection.TypeText(rez[2])    
            for j in rez[1].split(";"):
                if j.upper()=="AL":
                    wordapp.Selection.ParagraphFormat.Alignment = win32com.client.constants.wdAlignParagraphLeft
                elif j.upper()=="AR":
                    wordapp.Selection.ParagraphFormat.Alignment = win32com.client.constants.wdAlignParagraphRight
                elif j.upper()=="AC":
                    wordapp.Selection.ParagraphFormat.Alignment = win32com.client.constants.wdAlignParagraphCenter
                elif j.upper()=="B":
                    wordapp.Selection.MoveLeft(Unit=win32com.client.constants.wdCharacter, Count=len(rez[2]), Extend=win32com.client.constants.wdExtend)
                    wordapp.Selection.Font.Bold = True
                    wordapp.Selection.Collapse(Direction=win32com.client.constants.wdCollapseEnd)
                elif j.upper()=="U":
                    wordapp.Selection.MoveLeft(Unit=win32com.client.constants.wdCharacter, Count=len(rez[2]), Extend=win32com.client.constants.wdExtend)
                    wordapp.Selection.Font.Underline = win32com.client.constants.wdUnderlineSingle
                    wordapp.Selection.Collapse(Direction=win32com.client.constants.wdCollapseEnd)

        else:
            wordapp.Selection.TypeText(i)    
        
        wordapp.Selection.MoveRight(Unit=win32com.client.constants.wdCell)        

        
def word_find_tbl_by_descr(worddoc, descr, passnum):
    #найти в документе word таблицу с заданным описанием, пропустив passnum таких таблиц
    for tab in worddoc.Tables:
        if tab.Descr == descr:
            if passnum > 0:
                passnum = passnum - 1
            else:
                return tab
    raise Exception("Table not found. Table name=%s, num=%d" % (descr, passnum))
        

def word_table_fill(data, fdata, wordapp, worddoc, tabname, text_next, split_header=False):
    #Заполняем вордовскую таблицу (11 колоночную) у которой 1 строка и первая колонка - заголовки
    #data - массив массивов для вывода в табличку
    #fdata - массив форматов для вывода данных
    #wordapp - приложение word
    #worddoc - открытый doc файл
    #tabname - имя таблицы (описание)
    #text_next - "Продолжение таблицы 2"
    #split_header - разбивать ячейку 1,2 и удалять пустые столбцы?

    data_len = len(data[0])
    data_num = len(data)
    empty_col = 0
    #кол-во табличек чтобы вместить все данные
    numtables = ceil(data_len / 10)
    for tabnum in range(numtables):
        tab = word_find_tbl_by_descr(worddoc, tabname, tabnum)
        for i in range(2,12):
            if i-1+(tabnum)*10 > data_len:
                empty_col+=1
                for j in range(data_num):
                    tab.Cell(j+2,i).Range.Text = ""
            else:
                for j in range(data_num):
                    tab.Cell(j+2,i).Range.Text = fdata[j] % data[j][i-2+tabnum*10]
             
        #последнюю таблицу не копируем
        if tabnum==numtables-1:
            break
            
        tab.Select()
        wordapp.Selection.Copy()
        wordapp.Selection.MoveRight()
        wordapp.Selection.TypeText(text_next)
        wordapp.Selection.Paste()

    if split_header and empty_col:
        deletecol = 11 - empty_col
        #Удаляем пустые ячейки из последней таблички и форматируем ее по ширине.
        tab.Cell(1, 2).Split(NumRows=1, NumColumns=10)
        for i in range(11,deletecol,-1):
            tab.Columns(i).Delete()
        if deletecol>3:
            tab.Cell(1, 2).Merge(MergeTo=tab.Cell(Row=1, Column=deletecol))
        tab.AutoFitBehavior (win32com.client.constants.wdAutoFitContent)
        tab.AutoFitBehavior (win32com.client.constants.wdAutoFitWindow)

